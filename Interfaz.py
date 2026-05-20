from textwrap import wrap
import tkinter as tk
from tkinter import messagebox, filedialog
from motor_inferencia import (recomendar_por_datos, responder_modo_libre, es_pregunta_de_guia, generar_guia_metodologia, responder_seguimiento, es_consulta_de_seguimiento)
from base_conocimiento import NOMBRES_CONDICIONES


# =========================
# CONFIGURACIÓN GENERAL
# =========================n

ventana = tk.Tk()
ventana.title("Sistema Experto - Metodologías de Desarrollo")
ventana.attributes("-zoomed", True)
ventana.configure(bg="#061527")

COLOR_FONDO = "#061527"
COLOR_PANEL = "#0b1f36"
COLOR_PANEL_2 = "#102a46"
COLOR_AZUL = "#2563eb"
COLOR_VERDE = "#22c55e"
COLOR_TEXTO = "#f8fafc"
COLOR_TEXTO_SEC = "#cbd5e1"
COLOR_BURBUJA_USER = "#0b5ed7"
COLOR_BURBUJA_BOT = "#132f4c"

COLORES_METODOLOGIAS = {
    "Scrum": "#2563eb",        # azul
    "Kanban": "#38bdf8",       # celeste
    "Cascada": "#64748b",      # gris azulado
    "XP": "#a855f7",           # morado
    "Lean": "#22c55e",         # verde
    "Espiral": "#f97316",      # naranja
    "Incremental": "#14b8a6",  # turquesa
    "RAD": "#eab308",          # amarillo
    "SAFe": "#ef4444"          # rojo
}

# =========================
# FUNCIONES GENERALES
# =========================

def limpiar_ventana():
    for widget in ventana.winfo_children():
        widget.destroy()

def exportar_resultado_word(texto_resultado):

    if not texto_resultado.strip():
        messagebox.showwarning("Sin resultado", "No hay ningún resultado para exportar.")
        return

    ruta = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documento de Word", "*.docx")],
        title="Guardar resultado"
    )

    if not ruta:
        return

    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        titulo = doc.add_paragraph()
        run = titulo.add_run("Resultado del Sistema Experto")
        run.bold = True
        run.font.size = Pt(18)

        doc.add_paragraph()

        for linea in texto_resultado.split("\n"):
            p = doc.add_paragraph()
            r = p.add_run(linea)
            r.font.size = Pt(12)

        doc.save(ruta)

        messagebox.showinfo("Exportación exitosa", "El resultado se exportó correctamente.")

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo exportar el documento:\n{e}")

def exportar_reporte_word(descripcion, resultado):
    if resultado is None:
        messagebox.showwarning("Sin resultado", "Primero genera una recomendación para poder exportar el reporte.")
        return

    ruta = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documento de Word", "*.docx")],
        title="Guardar reporte del sistema experto"
    )

    if not ruta:
        return

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from datetime import datetime

        doc = Document()

        # =========================
        # PORTADA / TÍTULO
        # =========================

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("Reporte del Sistema Experto")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(31, 78, 121)

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitulo.add_run("Recomendación de Metodologías de Desarrollo de Software")
        run.font.size = Pt(14)

        fecha = doc.add_paragraph()
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fecha.add_run(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        run.font.size = Pt(11)

        doc.add_paragraph()

        # =========================
        # 1. DATOS GENERALES
        # =========================

        doc.add_heading("1. Datos generales del análisis", level=1)

        p = doc.add_paragraph()
        p.add_run("Objetivo del reporte: ").bold = True
        p.add_run(
            "Presentar la recomendación generada por el sistema experto con base en las "
            "condiciones detectadas del proyecto descrito por el usuario."
        )

        p = doc.add_paragraph()
        p.add_run("Metodología recomendada: ").bold = True
        p.add_run(resultado["principal"])

        p = doc.add_paragraph()
        p.add_run("Alternativa recomendada: ").bold = True
        p.add_run(resultado["alternativa"])

        p = doc.add_paragraph()
        p.add_run("Nivel de confianza: ").bold = True
        p.add_run(resultado.get("confianza_nivel", "No calculado"))

        doc.add_paragraph(resultado.get("confianza_detalle", ""))

        # =========================
        # 2. DESCRIPCIÓN
        # =========================

        doc.add_heading("2. Descripción del proyecto analizado", level=1)

        if descripcion.strip():
            doc.add_paragraph(descripcion)
        else:
            doc.add_paragraph("No se registró una descripción del proyecto.")

        # =========================
        # 3. CONDICIONES DETECTADAS
        # =========================

        doc.add_heading("3. Condiciones detectadas", level=1)

        condiciones = resultado.get("condiciones_detectadas", [])

        if condiciones:
            for condicion in condiciones:
                nombre_bonito = NOMBRES_CONDICIONES.get(
                    condicion,
                    condicion.replace("_", " ")
                )
                doc.add_paragraph(nombre_bonito, style="List Bullet")
        else:
            doc.add_paragraph(
                "No se registraron condiciones detectadas automáticamente. "
                "Esto puede ocurrir cuando el análisis proviene del modo guiado."
            )

        # =========================
        # 4. PUNTAJES
        # =========================

        doc.add_heading("4. Puntajes por metodología", level=1)

        tabla = doc.add_table(rows=1, cols=2)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.style = "Table Grid"

        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Metodología"
        hdr_cells[1].text = "Puntaje"

        for metodologia, puntaje in resultado["puntajes"]:
            row_cells = tabla.add_row().cells
            row_cells[0].text = metodologia
            row_cells[1].text = str(puntaje)

        doc.add_paragraph()

        # =========================
        # 5. JUSTIFICACIÓN
        # =========================

        doc.add_heading("5. Justificación de la recomendación", level=1)
        doc.add_paragraph(resultado["justificacion_principal"])

        doc.add_heading("6. Análisis de la alternativa", level=1)
        doc.add_paragraph(resultado["justificacion_alternativa"])

        # =========================
        # 7. GUÍA DE APLICACIÓN
        # =========================

        doc.add_heading("7. Guía práctica de aplicación", level=1)

        try:
            guia = generar_guia_metodologia(resultado["principal"])
            doc.add_paragraph(guia)
        except:
            doc.add_paragraph(
                "No se encontró una guía específica para la metodología recomendada."
            )

        # =========================
        # 8. CONCLUSIÓN
        # =========================

        doc.add_heading("8. Conclusión", level=1)

        conclusion = (
            f"Con base en el análisis realizado, el sistema experto recomienda utilizar "
            f"{resultado['principal']} como metodología principal, debido a que presenta "
            f"mayor compatibilidad con las características del proyecto. "
            f"Como alternativa, se sugiere considerar {resultado['alternativa']}, ya que "
            f"también obtuvo un puntaje relevante dentro del análisis."
        )

        doc.add_paragraph(conclusion)

        # =========================
        # FORMATO GENERAL
        # =========================

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                if run.font.size is None:
                    run.font.size = Pt(12)

        doc.save(ruta)

        messagebox.showinfo("Reporte exportado", "El reporte se exportó correctamente a Word.")

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo exportar el reporte:\n{e}")

def mostrar_grafica_puntajes(puntajes):
    ventana_grafica = tk.Toplevel(ventana)
    ventana_grafica.title("Gráfica de puntajes")

    ancho = 780
    alto = 520

    pantalla_ancho = ventana_grafica.winfo_screenwidth()
    pantalla_alto = ventana_grafica.winfo_screenheight()

    x = int((pantalla_ancho / 2) - (ancho / 2))
    y = int((pantalla_alto / 2) - (alto / 2))

    ventana_grafica.geometry(f"{ancho}x{alto}+{x}+{y}")
    ventana_grafica.configure(bg=COLOR_FONDO)

    tk.Label(
        ventana_grafica,
        text="Puntajes por metodología",
        font=("Arial", 18, "bold"),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO
    ).pack(pady=(20, 5))

    tk.Label(
        ventana_grafica,
        text="Comparación visual de compatibilidad según las condiciones del proyecto.",
        font=("Arial", 11),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO_SEC
    ).pack(pady=(0, 15))

    canvas = tk.Canvas(
        ventana_grafica,
        width=700,
        height=360,
        bg=COLOR_PANEL,
        highlightthickness=0
    )
    canvas.pack(padx=30, pady=10)

    if not puntajes:
        canvas.create_text(
            350, 180,
            text="No hay puntajes disponibles.",
            fill=COLOR_TEXTO,
            font=("Arial", 13, "bold")
        )
        return

    max_puntaje = max(puntaje for _, puntaje in puntajes)

    if max_puntaje == 0:
        max_puntaje = 1

    margen_izq = 150
    margen_der = 60
    ancho_max_barra = 700 - margen_izq - margen_der

    y_inicial = 35
    separacion = 35
    alto_barra = 20

    colores = {
        "Scrum": "#2563eb",
        "Kanban": "#38bdf8",
        "Cascada": "#64748b",
        "XP": "#a855f7",
        "Lean": "#22c55e",
        "Espiral": "#f97316",
        "Incremental": "#14b8a6",
        "RAD": "#eab308",
        "SAFe": "#ef4444"
    }

    for i, (metodologia, puntaje) in enumerate(puntajes):
        y = y_inicial + (i * separacion)

        color = colores.get(metodologia, "#2563eb")

        ancho_barra = int((puntaje / max_puntaje) * ancho_max_barra)

        # Nombre de metodología
        canvas.create_text(
            20, y + 10,
            text=metodologia,
            anchor="w",
            fill=COLOR_TEXTO,
            font=("Arial", 11, "bold")
        )

        # Fondo de barra
        canvas.create_rectangle(
            margen_izq,
            y,
            margen_izq + ancho_max_barra,
            y + alto_barra,
            fill="#1e3a5f",
            outline=""
        )

        # Barra real
        canvas.create_rectangle(
            margen_izq,
            y,
            margen_izq + ancho_barra,
            y + alto_barra,
            fill=color,
            outline=""
        )

        # Puntaje
        canvas.create_text(
            margen_izq + ancho_barra + 10,
            y + 10,
            text=str(puntaje),
            anchor="w",
            fill=COLOR_TEXTO,
            font=("Arial", 10, "bold")
        )

    tk.Button(
        ventana_grafica,
        text="Cerrar",
        font=("Arial", 11, "bold"),
        bg=COLOR_AZUL,
        fg="white",
        relief="flat",
        padx=20,
        pady=8,
        command=ventana_grafica.destroy
    ).pack(pady=15)

def crear_layout_base(modo_activo="libre"):
    limpiar_ventana()

    contenedor = tk.Frame(ventana, bg=COLOR_FONDO)
    contenedor.pack(fill="both", expand=True)

    sidebar = tk.Frame(contenedor, bg="#071b2f", width=580)
    sidebar.pack(side="left", fill="y", padx=(15, 5), pady=15)
    sidebar.pack_propagate(False)

    main = tk.Frame(contenedor, bg=COLOR_FONDO)
    main.pack(side="right", fill="both", expand=True, padx=(5, 15), pady=15)

    # Logo / título sidebar
    tk.Label(
        sidebar,
        text="🤖",
        font=("Arial", 44),
        bg="#071b2f",
        fg=COLOR_TEXTO
    ).pack(pady=(35, 5))

    tk.Label(
        sidebar,
        text="Sistema Experto",
        font=("Arial", 20, "bold"),
        bg="#071b2f",
        fg=COLOR_TEXTO
    ).pack()

    tk.Label(
        sidebar,
        text="Metodologías de Desarrollo\n de Software",
        font=("Arial", 12),
        bg="#071b2f",
        fg=COLOR_TEXTO_SEC
    ).pack(pady=(5, 35))

    tk.Label(
        sidebar,
        text="MODOS",
        font=("Arial", 11, "bold"),
        bg="#071b2f",
        fg="#94a3b8"
    ).pack(anchor="w", padx=25)

    color_libre = COLOR_AZUL if modo_activo == "libre" else COLOR_PANEL
    color_guiado = COLOR_AZUL if modo_activo == "guiado" else COLOR_PANEL

    tk.Button(
        sidebar,
        text="💬  Modo libre",
        font=("Arial", 12, "bold"),
        bg=color_libre,
        fg="white",
        activebackground=COLOR_AZUL,
        activeforeground="white",
        relief="flat",
        anchor="w",
        padx=20,
        pady=12,
        command=pantalla_libre_chat
    ).pack(fill="x", padx=25, pady=(15, 8))

    tk.Button(
        sidebar,
        text="📝  Modo guiado",
        font=("Arial", 12, "bold"),
        bg=color_guiado,
        fg="white",
        activebackground=COLOR_AZUL,
        activeforeground="white",
        relief="flat",
        anchor="w",
        padx=20,
        pady=12,
        command=pantalla_guiada
    ).pack(fill="x", padx=25, pady=8)

    separador = tk.Frame(sidebar, bg="#1e3a5f", height=1)
    separador.pack(fill="x", padx=25, pady=25)

    tk.Label(
        sidebar,
        text="ACERCA DEL SISTEMA",
        font=("Arial", 11, "bold"),
        bg="#071b2f",
        fg="#94a3b8"
    ).pack(anchor="w", padx=25)

    tk.Label(
        sidebar,
        text="Este sistema experto ayuda a elegir\n"
             "una metodología de desarrollo de\n"
             "software de acuerdo con las\n"
             "características del proyecto.",
        font=("Arial", 11),
        bg="#071b2f",
        fg=COLOR_TEXTO_SEC,
        justify="left"
    ).pack(anchor="w", padx=25, pady=(10, 30))

    tarjeta = tk.Frame(sidebar, bg=COLOR_PANEL, padx=15, pady=15)
    tarjeta.pack(fill="x", padx=25, pady=(15, 10))

    tk.Label(
        tarjeta,
        text="🧠 Inteligencia basada en reglas",
        font=("Arial", 11, "bold"),
        bg=COLOR_PANEL,
        fg=COLOR_TEXTO,
        justify="left"
    ).pack(anchor="w")

    tk.Label(
        tarjeta,
        text="Analiza condiciones, aplica reglas\n"
             "y genera una recomendación\n"
             "justificada.",
        font=("Arial", 10),
        bg=COLOR_PANEL,
        fg=COLOR_TEXTO_SEC,
        justify="left"
    ).pack(anchor="w", pady=(10, 0))

    return main

def crear_grafica_en_frame(parent, puntajes):
    canvas = tk.Canvas(
        parent,
        width=900,
        height=390,
        bg=COLOR_PANEL,
        highlightthickness=0
    )
    canvas.pack(anchor="w", pady=(10, 15))

    if not puntajes:
        canvas.create_text(
            325, 150,
            text="No hay puntajes disponibles.",
            fill=COLOR_TEXTO,
            font=("Arial", 12, "bold")
        )
        return canvas

    max_puntaje = max(puntaje for _, puntaje in puntajes)

    if max_puntaje == 0:
        max_puntaje = 1

    margen_izq = 150
    ancho_max = 600
    y_inicial = 35
    separacion = 34
    alto_barra = 18

    canvas.create_text(
        20,
        10,
        text="📊 Comparación de puntajes",
        anchor="w",
        fill=COLOR_TEXTO,
        font=("Arial", 12, "bold")
    )

    for i, (metodologia, puntaje) in enumerate(puntajes):
        y = y_inicial + (i * separacion)
        color = COLORES_METODOLOGIAS.get(metodologia, COLOR_AZUL)

        ancho_barra = int((puntaje / max_puntaje) * ancho_max)

        canvas.create_text(
            20,
            y + 9,
            text=metodologia,
            anchor="w",
            fill=COLOR_TEXTO,
            font=("Arial", 10, "bold")
        )

        canvas.create_rectangle(
            margen_izq,
            y,
            margen_izq + ancho_max,
            y + alto_barra,
            fill="#1e3a5f",
            outline=""
        )

        canvas.create_rectangle(
            margen_izq,
            y,
            margen_izq + ancho_barra,
            y + alto_barra,
            fill=color,
            outline=""
        )

        canvas.create_text(
            margen_izq + ancho_barra + 8,
            y + 9,
            text=str(puntaje),
            anchor="w",
            fill=COLOR_TEXTO,
            font=("Arial", 10, "bold")
        )

    return canvas

# =========================
# FORMATEAR RESULTADO
# =========================

def formatear_resultado(resultado):
    texto_puntajes = ""
    for metodologia, puntaje in resultado["puntajes"]:
        texto_puntajes += f"{metodologia}: {puntaje}\n"

    texto_condiciones = ""

    if "condiciones_detectadas" in resultado:
        condiciones = resultado.get("condiciones_detectadas", [])

        if condiciones:
            texto_condiciones += "\n✅ Condiciones detectadas:\n"
            for condicion in condiciones:
                nombre_bonito = NOMBRES_CONDICIONES.get(
                    condicion,
                    condicion.replace("_", " ")
                )
                texto_condiciones += f"• {nombre_bonito}\n"

    texto = (
        f"🏆 Metodología recomendada: {resultado['principal']}\n"
        f"🔁 Alternativa recomendada: {resultado['alternativa']}\n\n"
        f"📌 Nivel de confianza: {resultado.get('confianza_nivel', 'No calculado')}\n"
        f"{resultado.get('confianza_detalle', '')}\n\n"
        f"📊 Puntajes:\n{texto_puntajes}"
        f"{texto_condiciones}\n"
        f"📌 Justificación:\n{resultado['justificacion_principal']}\n\n"
        f"⚖️ Alternativa:\n{resultado['justificacion_alternativa']}"
    )

    return texto

def formatear_resultado_resumido(resultado):
    texto_condiciones = ""

    if "condiciones_detectadas" in resultado:
        condiciones = resultado.get("condiciones_detectadas", [])

        if condiciones:
            texto_condiciones += "Detecté principalmente estas condiciones:\n"

            for condicion in condiciones[:5]:
                nombre_bonito = NOMBRES_CONDICIONES.get(
                    condicion,
                    condicion.replace("_", " ")
                )
                texto_condiciones += f"• {nombre_bonito}\n"

    texto = (
        f"🏆 Metodología recomendada: {resultado['principal']}\n"
        f"🔁 Alternativa recomendada: {resultado['alternativa']}\n"
        f"📌 Nivel de confianza: {resultado.get('confianza_nivel', 'No calculado')}\n\n"
        f"{texto_condiciones}\n"
        f"Con base en estas características, {resultado['principal']} parece ser la opción más adecuada para iniciar el proyecto.\n\n"
        f"Puedes preguntarme cosas como:\n"
        f"• ¿Por qué me recomendaste {resultado['principal']}?\n"
        f"• ¿Cómo puedo usar {resultado['principal']}?\n"
        f"• ¿Qué ventajas tiene?\n"
        f"• ¿Qué desventajas tiene?\n"
        f"• ¿Cuál es la diferencia con {resultado['alternativa']}?\n"
        f"• Dame un ejemplo aplicado a mi proyecto."
    )

    return texto

def agregar_mensaje(chat, autor, mensaje, tipo="bot"):
    chat.config(state="normal")

    if tipo == "user":
        chat.insert("end", "\nTú:\n", "user_title")
        chat.insert("end", mensaje + "\n", "user_msg")
    else:
        chat.insert("end", "\nSistema Experto:\n", "bot_title")
        chat.insert("end", mensaje + "\n", "bot_msg")

    chat.config(state="disabled")
    chat.see("end")


# =========================
# MODO LIBRE TIPO CHAT
# =========================

def pantalla_libre_chat():
    main = crear_layout_base("libre")

    header = tk.Frame(main, bg=COLOR_FONDO)
    header.pack(fill="x", padx=40, pady=(35, 10))

    tk.Label(
        header,
        text="¡Hola! Describe tu Proyecto",
        font=("Arial", 26, "bold"),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO
    ).pack(anchor="w")

    tk.Label(
        header,
        text="Cuéntame las características principales de tu proyecto y te recomendaré la metodología más adecuada.",
        font=("Arial", 12),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO_SEC
    ).pack(anchor="w", pady=(8, 0))


    # Chat con burbujas
    chat_outer = tk.Frame(main, bg=COLOR_FONDO)
    chat_outer.pack(fill="both", expand=True, padx=40, pady=20)

    canvas = tk.Canvas(
        chat_outer,
        bg=COLOR_FONDO,
        highlightthickness=0
    )
    canvas.pack(side="left", fill="both", expand=True)

    #scrollbar = tk.Scrollbar(chat_outer, command=canvas.yview)
    #scrollbar.pack(side="right", fill="y")

    canvas.configure(yscrollcommand=None)

    def mover_scroll(event):
        nonlocal auto_scroll
        auto_scroll = False
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", mover_scroll))
    canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))

    chat_frame = tk.Frame(canvas, bg=COLOR_FONDO)

    chat_window = canvas.create_window(
        (0, 0),
        window=chat_frame,
        anchor="nw"
)

    auto_scroll = True

    def ajustar_ancho_canvas(event):
        canvas.itemconfig(chat_window, width=event.width)

    canvas.bind("<Configure>", ajustar_ancho_canvas)

    def actualizar_scroll(event=None):
        canvas.configure(scrollregion=canvas.bbox("all"))
        if auto_scroll:
            canvas.yview_moveto(1)
    chat_frame.bind("<Configure>", actualizar_scroll)

    def agregar_burbuja(mensaje, tipo="bot"):
        fila = tk.Frame(chat_frame, bg=COLOR_FONDO)
        fila.pack(fill="x", pady=8, padx=10)

        if tipo == "user":
            contenedor_burbuja = tk.Frame(fila, bg=COLOR_FONDO)
            contenedor_burbuja.pack(side="right", anchor="e")

            burbuja = tk.Frame(contenedor_burbuja, bg="#0b5ed7", padx=15, pady=12)
            burbuja.pack(anchor="e")

            color_texto = "white"
            titulo = "Tú"
            ancho_texto = 520
        else:   
            contenedor_burbuja = tk.Frame(fila, bg=COLOR_FONDO)
            contenedor_burbuja.pack(side="left", anchor="w")

            burbuja = tk.Frame(contenedor_burbuja, bg=COLOR_PANEL_2, padx=15, pady=12)
            burbuja.pack(anchor="w")

            color_texto = COLOR_TEXTO
            titulo = "Sistema Experto"
            ancho_texto = 650

        tk.Label(
            burbuja,
            text=titulo,
            font=("Arial", 10, "bold"),
            bg=burbuja["bg"],
            fg="#93c5fd" if tipo == "bot" else "#bfdbfe"
        ).pack(anchor="w")

        tk.Label(
            burbuja,
            text=mensaje,
            font=("Arial", 11),
            bg=burbuja["bg"],
            fg=color_texto,
            wraplength=ancho_texto,
            justify="left"
        ).pack(anchor="w", pady=(5, 0))

    def agregar_resultado_con_grafica(resultado):
        fila = tk.Frame(chat_frame, bg=COLOR_FONDO)
        fila.pack(fill="x", pady=8, padx=10)

        color_metodologia = COLORES_METODOLOGIAS.get(resultado["principal"], COLOR_AZUL)

        contenedor_burbuja = tk.Frame(fila, bg=COLOR_FONDO)
        contenedor_burbuja.pack(side="left", anchor="w")

        burbuja = tk.Frame(
            contenedor_burbuja,
            bg=COLOR_PANEL_2,
            padx=18,
            pady=15)
        burbuja.pack(anchor="w")

        tk.Frame(
            burbuja,
            bg=color_metodologia,
            height=5).pack(fill="x", pady=(0, 12))    

        tk.Label(
            burbuja,
            text="Sistema Experto",
            font=("Arial", 10, "bold"),
            bg=COLOR_PANEL_2,
            fg="#93c5fd"
        ).pack(anchor="w")

        tk.Label(
            burbuja,
            text="🏆 Metodología recomendada",
            font=("Arial", 12, "bold"),
            bg=COLOR_PANEL_2,
            fg=COLOR_TEXTO
        ).pack(anchor="w", pady=(5, 2))

        tk.Label(
            burbuja,
            text=resultado["principal"],
            font=("Arial", 20, "bold"),
            bg=COLOR_PANEL_2,
            fg=color_metodologia
        ).pack(anchor="w", pady=(0, 8))

        color_alternativa = COLORES_METODOLOGIAS.get(resultado["alternativa"], COLOR_TEXTO_SEC)

        tk.Label(
            burbuja,
            text=f"🔁 Alternativa recomendada: {resultado['alternativa']}",
            font=("Arial", 11, "bold"),
            bg=COLOR_PANEL_2,
            fg=color_alternativa
        ).pack(anchor="w", pady=(0, 10))

        tk.Label(
            burbuja,
            text=f"📌 Nivel de confianza: {resultado.get('confianza_nivel', 'No calculado')}",
            font=("Arial", 11),
            bg=COLOR_PANEL_2,
            fg=COLOR_TEXTO,
            wraplength=950,
            justify="left"
        ).pack(anchor="w", pady=(0, 8))

        crear_grafica_en_frame(burbuja, resultado["puntajes"])

        texto_condiciones = ""

        if "condiciones_detectadas" in resultado:
            condiciones = resultado.get("condiciones_detectadas", [])

            if condiciones:
                texto_condiciones += "✅ Condiciones detectadas:\n"
                for condicion in condiciones:
                    nombre_bonito = NOMBRES_CONDICIONES.get(
                        condicion,
                        condicion.replace("_", " ")
                    )
                    texto_condiciones += f"• {nombre_bonito}\n"
                texto_condiciones += "\n"

        cuerpo = (
            f"{texto_condiciones}"
            f"Con base en estas características, {resultado['principal']} parece ser la opción más adecuada para iniciar el proyecto.\n\n"
            f"Puedes preguntarme cosas como:\n"
            f"• ¿Por qué me recomendaste {resultado['principal']}?\n"
            f"• ¿Cómo puedo usar {resultado['principal']}?\n"
            f"• ¿Qué ventajas tiene?\n"
            f"• ¿Qué desventajas tiene?\n"
            f"• ¿Cuál es la diferencia con {resultado['alternativa']}?\n"
            f"• Dame un ejemplo aplicado a mi proyecto."
            )

        tk.Label(
            burbuja,
            text=cuerpo,
            font=("Arial", 11),
            bg=COLOR_PANEL_2,
            fg=COLOR_TEXTO,
            wraplength=950,
            justify="left"
        ).pack(anchor="w", pady=(5, 0))

    actualizar_scroll()
    agregar_burbuja(
        "🤖\nDime las características de tu proyecto para recomendarte la mejor metodología.",
        "bot"
    )

    def agregar_burbuja_animada(mensaje, tipo="bot", velocidad=8):
        fila = tk.Frame(chat_frame, bg=COLOR_FONDO)
        fila.pack(fill="x", pady=8, padx=10)

        ancho_chat = canvas.winfo_width()

        if ancho_chat <= 1:
            ancho_chat = 1000

        if tipo == "user":
            contenedor_burbuja = tk.Frame(fila, bg=COLOR_FONDO)
            contenedor_burbuja.pack(side="right", anchor="e")

            burbuja = tk.Frame(contenedor_burbuja, bg="#0b5ed7", padx=15, pady=12)
            burbuja.pack(anchor="e")

            color_texto = "white"
            titulo = "Tú"
            ancho_texto = int(ancho_chat * 0.55)
        else:
            contenedor_burbuja = tk.Frame(fila, bg=COLOR_FONDO)
            contenedor_burbuja.pack(side="left", anchor="w", fill="x", expand=True)

            burbuja = tk.Frame(contenedor_burbuja, bg=COLOR_PANEL_2, padx=15, pady=12)
            burbuja.pack(anchor="w")

            color_texto = COLOR_TEXTO
            titulo = "Sistema Experto"
            ancho_texto = int(ancho_chat * 0.70)

        tk.Label(
            burbuja,
            text=titulo,
            font=("Arial", 10, "bold"),
            bg=burbuja["bg"],
            fg="#93c5fd" if tipo == "bot" else "#bfdbfe"
        ).pack(anchor="w")

        etiqueta = tk.Label(
            burbuja,
            text="",
            font=("Arial", 11),
            bg=burbuja["bg"],
            fg=color_texto,
            wraplength=ancho_texto,
            justify="left"
        )
        etiqueta.pack(anchor="w", pady=(5, 0))

        def escribir(i=0):
            if i <= len(mensaje):
                etiqueta.config(text=mensaje[:i])
                actualizar_scroll()
                ventana.after(velocidad, escribir, i + 1)

        escribir()
        
    # Entrada inferior
    entrada_frame = tk.Frame(main, bg=COLOR_PANEL_2, padx=15, pady=12)
    entrada_frame.pack(fill="x", padx=40, pady=(0, 30))

    entrada = tk.Text(
        entrada_frame,
        height=3,
        font=("Arial", 12),
        bg=COLOR_PANEL_2,
        fg=COLOR_TEXTO,
        insertbackground="white",
        relief="flat",
        wrap="word"
    )
    entrada.pack(side="left", fill="x", expand=True, padx=(0, 10))

    contexto_chat = {
    "metodologia": None,
    "alternativa": None,
    "texto": "",
    "puntajes": [],
    "descripcion": "",
    "resultado": None,
    "seguimiento_iniciado": False
    }

    def enviar():
        try:
            texto = entrada.get("1.0", tk.END).strip()

            if texto == "":
                messagebox.showwarning("Falta información", "Escribe una descripción del proyecto.")
                return

            entrada.delete("1.0", tk.END)

            agregar_burbuja(texto, "user")

            # Preguntas de seguimiento: roles, ventajas, documentos, diferencia, etc.
            if contexto_chat["resultado"] is not None and es_consulta_de_seguimiento(texto):
                print("CONTEXTO USADO:", contexto_chat["metodologia"], contexto_chat["alternativa"])

                respuesta_seguimiento = responder_seguimiento(
                    texto,
                    contexto_chat["metodologia"],
                    contexto_chat["alternativa"],
                    contexto_chat["resultado"],
                    contexto_chat["descripcion"],
                    contexto_chat["seguimiento_iniciado"]
                )

                if respuesta_seguimiento:
                    agregar_burbuja_animada(respuesta_seguimiento, "bot")
                    contexto_chat["seguimiento_iniciado"] = True    
                    return
                

            # Pregunta específica: cómo usar la metodología recomendada
            if es_pregunta_de_guia(texto):
                if contexto_chat["metodologia"] is None:
                    agregar_burbuja_animada(
                        "Aún no tengo una metodología recomendada. Primero descríbeme tu proyecto y después puedo explicarte cómo aplicarla.",
                        "bot"
                    )
                    return
                else:
                    guia = generar_guia_metodologia(contexto_chat["metodologia"])
                    agregar_burbuja_animada(guia, "bot")
                    return

            # Si no es seguimiento, se analiza como nueva descripción de proyecto
            respuesta_sistema = responder_modo_libre(texto)

            if respuesta_sistema["tipo"] == "mensaje":
                agregar_burbuja_animada(respuesta_sistema["respuesta"], "bot")
                return
            resultado = respuesta_sistema["resultado"]

            respuesta_completa = formatear_resultado(resultado)

            contexto_chat["metodologia"] = resultado["principal"]
            contexto_chat["alternativa"] = resultado["alternativa"]
            contexto_chat["texto"] = respuesta_completa
            contexto_chat["puntajes"] = resultado["puntajes"]
            contexto_chat["descripcion"] = texto
            contexto_chat["resultado"] = resultado
            contexto_chat["seguimiento_iniciado"] = False

            print("CONTEXTO GUARDADO:", contexto_chat["metodologia"], contexto_chat["alternativa"])

            agregar_resultado_con_grafica(resultado)

        except Exception as e:
            print("ERROR EN MODO LIBRE:", e)
            messagebox.showerror("Error", f"Ocurrió un error en el modo libre:\n{e}")
    
#---------------------------------------------    
    boton_enviar = tk.Button(
        entrada_frame,
        text="➤",
        font=("Arial", 16, "bold"),
        bg="#ff5a3c",
        fg="white",
        activebackground="#ef4444",
        activeforeground="white",
        relief="flat",
        width=5,
        height=2,
        command=enviar
    )
    boton_enviar.pack(side="right", padx=(10, 10))

    boton_word = tk.Button(
        entrada_frame,
        text="Exportar",
        font=("Arial", 16, "bold"),
        bg="#ff5a3c",
        fg="white",
        activebackground="#ef4444",
        activeforeground="white",
        relief="flat",
        width=10,
        height=2,
        command=lambda: exportar_reporte_word(
            contexto_chat["descripcion"],
            contexto_chat["resultado"]
        )
    )
    boton_word.pack(side="right", padx=(10, 0))

    entrada.bind("<Control-Return>", lambda event: enviar())

# =========================
# MODO GUIADO
# =========================

def pantalla_guiada():
    main = crear_layout_base("guiado")

    tk.Label(
        main,
        text="Modo guiado",
        font=("Arial", 24, "bold"),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO
    ).pack(pady=(30, 5))

    tk.Label(
        main,
        text="Responde las preguntas para obtener una recomendación.",
        font=("Arial", 12),
        bg=COLOR_FONDO,
        fg=COLOR_TEXTO_SEC
    ).pack(pady=(0, 20))

    preguntas = {
    "requisitos_cambian": "¿Las funciones o requisitos del proyecto podrían cambiar durante el desarrollo?",
    "cliente_participa": "¿El cliente o usuario final revisará avances y dará opiniones seguido?",
    "entregas_rapidas": "¿Se necesita entregar avances funcionales en poco tiempo?",
    "flujo_continuo": "¿El trabajo será constante, como soporte, mantenimiento o mejoras continuas?",
    "alta_calidad": "¿El sistema necesita muchas pruebas y código muy bien cuidado?",
    "proyecto_grande": "¿El proyecto es grande, complejo o importante para la organización?",
    "alto_riesgo": "¿Existe riesgo por tecnología nueva, dificultad técnica o falta de experiencia?",
    "documentacion": "¿Se necesita dejar mucha documentación formal del proyecto?",
    "tiempo_estricto": "¿La fecha de entrega es muy estricta y no se puede mover?",
    "optimizacion": "¿Se busca ahorrar tiempo, reducir desperdicios o usar mejor los recursos?",

    # Nuevas preguntas para metodologías agregadas
    "modularidad": "¿El proyecto puede dividirse en módulos, versiones o entregas parciales?",
    "prototipo": "¿Se necesita crear un prototipo rápido, demo o versión de prueba para validar la idea?",
    "varios_equipos": "¿Participan varios equipos, áreas o departamentos en el proyecto?",
    "escala_empresarial": "¿El proyecto pertenece a un contexto empresarial, corporativo o de gran escala?"
}

    respuestas = {}

    contenedor = tk.Frame(main, bg=COLOR_PANEL, padx=30, pady=20)
    contenedor.pack(padx=40, pady=10, fill="both")

    for clave, texto in preguntas.items():
        respuestas[clave] = tk.StringVar(value="")

        fila = tk.Frame(contenedor, bg=COLOR_PANEL)
        fila.pack(fill="x", pady=7)

        tk.Label(
            fila,
            text=texto,
            font=("Arial", 11),
            bg=COLOR_PANEL,
            fg=COLOR_TEXTO,
            wraplength=700,
            justify="left"
        ).pack(side="left", fill="x", expand=True)

        opciones = tk.Frame(fila, bg=COLOR_PANEL)
        opciones.pack(side="right")

        tk.Radiobutton(
            opciones,
            text="Sí",
            variable=respuestas[clave],
            value="si",
            bg=COLOR_PANEL,
            fg=COLOR_TEXTO,
            selectcolor=COLOR_PANEL_2,
            activebackground=COLOR_PANEL,
            activeforeground=COLOR_TEXTO,
            font=("Arial", 10)
        ).pack(side="left", padx=8)

        tk.Radiobutton(
            opciones,
            text="No",
            variable=respuestas[clave],
            value="no",
            bg=COLOR_PANEL,
            fg=COLOR_TEXTO,
            selectcolor=COLOR_PANEL_2,
            activebackground=COLOR_PANEL,
            activeforeground=COLOR_TEXTO,
            font=("Arial", 10)
        ).pack(side="left", padx=8)

    def evaluar_guiado():
        try:
            for clave, variable in respuestas.items():
                if variable.get() == "":
                    messagebox.showwarning("Faltan respuestas", "Por favor responde todas las preguntas.")
                    return

            datos = {clave: variable.get() for clave, variable in respuestas.items()}

            # Condiciones derivadas
            if datos["requisitos_cambian"] == "no":
                datos["requisitos_estables"] = "si"
            else:
                datos["requisitos_estables"] = "no"

            if datos["proyecto_grande"] == "no":
                datos["proyecto_pequeno"] = "si"
            else:
                datos["proyecto_pequeno"] = "no"

            resultado = recomendar_por_datos(datos)

            ventana_resultado = tk.Toplevel(ventana)
            ventana_resultado.title("Resultado del Sistema Experto")
            ventana_resultado.geometry("750x560")
            ventana_resultado.configure(bg=COLOR_FONDO)

            caja = tk.Text(
            ventana_resultado,
            wrap="word",
            font=("Arial", 11),
            bg=COLOR_PANEL,
            fg=COLOR_TEXTO,
            relief="flat",
            padx=20,
            pady=20
            )
            caja.pack(fill="both", expand=True, padx=20, pady=20)

            texto_resultado = formatear_resultado(resultado)

            caja.insert("1.0", texto_resultado)
            caja.config(state="disabled")

            botones_resultado = tk.Frame(ventana_resultado, bg=COLOR_FONDO)
            botones_resultado.pack(pady=(0, 15))

            tk.Button(
                botones_resultado,
                text="Ver gráfica",
                font=("Arial", 11, "bold"),
                bg="#60a5fa",
                fg="white",
                relief="flat",
                padx=15,
                pady=8,
                command=lambda: mostrar_grafica_puntajes(resultado["puntajes"])
            ).pack(side="left", padx=8)

            tk.Button(
                botones_resultado,
                text="Exportar a Word",
                font=("Arial", 11, "bold"),
                bg=COLOR_AZUL,
                fg="white",
                relief="flat",
                padx=15,
                pady=8,
                command=lambda: exportar_reporte_word("Analisis Realizado Mediante el Modo Guiado, con base a respuestas estrcturadas de si/no.", resultado)
            ).pack(side="left", padx=8)

        except Exception as e:
            
            print("ERROR EN MODO GUIADO:", e)
            messagebox.showerror("Error", f"Ocurrió un error en el modo guiado:\n{e}")

    def limpiar_respuestas():
        for variable in respuestas.values():
            variable.set("")

    botones = tk.Frame(main, bg=COLOR_FONDO)
    botones.pack(pady=20)

    tk.Button(
        botones,
        text="Obtener recomen-dación",
        font=("Arial", 12, "bold"),
        bg=COLOR_AZUL,
        fg="white",
        relief="flat",
        padx=20,
        pady=10,
        command=evaluar_guiado
    ).pack(side="left", padx=10)

    tk.Button(
        botones,
        text="Limpiar respuestas",
        font=("Arial", 12),
        bg=COLOR_PANEL_2,
        fg=COLOR_TEXTO,
        relief="flat",
        padx=20,
        pady=10,
        command=limpiar_respuestas
    ).pack(side="left", padx=10)


# EJECUTAR
pantalla_libre_chat()
ventana.mainloop()